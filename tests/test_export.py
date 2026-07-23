"""Unit and integration tests for document export capabilities (DOCX & PDF)."""

import json
import tempfile
import unittest
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import io

from app import create_app
from config import Config
from services.export_service import (
    ExportServiceError,
    export_resume_docx,
    export_resume_pdf,
    export_ats_report_docx,
    export_ats_report_pdf,
    export_jd_match_report_docx,
    export_jd_match_report_pdf,
)
from services.version_service import create_resume_version
from services.resume_store import save_resume


class TestExportService(unittest.TestCase):
    def setUp(self):
        self.sample_resume = {
            "personal": {
                "full_name": "Jane Doe",
                "email": "jane@example.com",
                "phone": "+1 (555) 019-2834",
                "linkedin": "linkedin.com/in/janedoe",
                "github": "github.com/janedoe",
            },
            "education": {
                "degree": "B.S. Computer Science",
                "college": "Stanford University",
                "graduation_year": "2023",
                "gpa": "3.9",
            },
            "experience": {
                "role": "Software Engineer",
                "company": "Tech Corp",
                "duration": "2023 - Present",
                "responsibilities": "- Developed high-throughput REST APIs.\n- Optimized DB queries.",
            },
            "projects": {
                "project_name": "AI Assistant",
                "description": "- Built LLM-powered tool.\n- Integrated SQLite database.",
                "technologies": "Python, Flask, OpenAI",
            },
            "skills": "Python, SQL, Flask, Docker, Git",
            "certifications": "AWS Certified Solutions Architect",
            "achievements": "First place in Hackathon 2023",
            "languages": "English (Native), Spanish (Fluent)",
        }

        self.sample_ats_data = {
            "score": 85,
            "analysis_type": "AI Assessment",
            "strengths": [
                "Includes standard contact information.",
                "Uses impactful action verbs in experience section.",
            ],
            "weaknesses": [
                "Could expand on measurable metrics in bullet points.",
            ],
            "suggestions": [
                "Quantify achievements with percentages and numerical metrics.",
            ],
        }

        self.sample_jd_match_data = {
            "match_percentage": 78,
            "analysis_type": "AI Assessment",
            "matching_skills": ["Python", "Flask", "SQL"],
            "missing_technical_skills": ["Kubernetes", "Redis"],
            "missing_soft_skills": ["Agile/Scrum Leadership"],
            "recommended_keywords": ["K8S", "IN-MEMORY CACHING"],
            "recommended_certifications": ["Certified Kubernetes Administrator"],
            "recommended_projects": [
                "Build a Redis-backed caching proxy for high traffic endpoints."
            ],
            "learning_roadmap": [
                "Step 1: Complete Kubernetes deployment tutorials.",
                "Step 2: Build sample microservice project.",
            ],
        }

        self.temp_dir = tempfile.TemporaryDirectory()
        self.db_path = Path(self.temp_dir.name) / "test_exports.db"

        class TestConfig(Config):
            TESTING = True
            SECRET_KEY = "test-secret"
            DATABASE_PATH = Path(self.temp_dir.name) / "test_exports.db"
            UPLOAD_FOLDER = Path(self.temp_dir.name) / "uploads"
            GENERATED_FOLDER = Path(self.temp_dir.name) / "uploads" / "generated"
            WTF_CSRF_ENABLED = False

        self.app = create_app(TestConfig)
        self.client = self.app.test_client()

    def tearDown(self):
        self.temp_dir.cleanup()

    # ---------------------------------------------------------------------------
    # 1. RESUME EXPORT TESTS
    # ---------------------------------------------------------------------------

    def test_export_resume_docx(self):
        docx_bytes = export_resume_docx(self.sample_resume)
        self.assertIsNotNone(docx_bytes)
        self.assertGreater(len(docx_bytes), 100)

        # Parse DOCX back
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("Jane Doe", full_text)
        self.assertIn("Stanford University", full_text)
        self.assertIn("Software Engineer", full_text)
        self.assertIn("Python, SQL, Flask", full_text)

    def test_export_resume_pdf(self):
        pdf_bytes = export_resume_pdf(self.sample_resume)
        self.assertIsNotNone(pdf_bytes)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))

        # Read PDF back
        reader = PdfReader(io.BytesIO(pdf_bytes))
        self.assertGreater(len(reader.pages), 0)
        page_text = reader.pages[0].extract_text()
        self.assertIn("Jane Doe", page_text)
        self.assertIn("Stanford University", page_text)

    # ---------------------------------------------------------------------------
    # 2. ATS REPORT EXPORT TESTS
    # ---------------------------------------------------------------------------

    def test_export_ats_report_docx(self):
        docx_bytes = export_ats_report_docx(self.sample_ats_data)
        self.assertIsNotNone(docx_bytes)

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("ATS Compatibility Analysis Report", full_text)
        self.assertIn("85 / 100", full_text)
        self.assertIn("Key Strengths", full_text)

    def test_export_ats_report_pdf(self):
        pdf_bytes = export_ats_report_pdf(self.sample_ats_data)
        self.assertIsNotNone(pdf_bytes)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))

        reader = PdfReader(io.BytesIO(pdf_bytes))
        page_text = reader.pages[0].extract_text()
        self.assertIn("ATS Compatibility Analysis Report", page_text)
        self.assertIn("85", page_text)

    # ---------------------------------------------------------------------------
    # 3. JD MATCH REPORT EXPORT TESTS
    # ---------------------------------------------------------------------------

    def test_export_jd_match_report_docx(self):
        docx_bytes = export_jd_match_report_docx(self.sample_jd_match_data)
        self.assertIsNotNone(docx_bytes)

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("Job Description Match Analysis Report", full_text)
        self.assertIn("78%", full_text)
        self.assertIn("Matching Skills", full_text)

    def test_export_jd_match_report_pdf(self):
        pdf_bytes = export_jd_match_report_pdf(self.sample_jd_match_data)
        self.assertIsNotNone(pdf_bytes)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))

        reader = PdfReader(io.BytesIO(pdf_bytes))
        page_text = reader.pages[0].extract_text()
        self.assertIn("Job Description Match Analysis Report", page_text)
        self.assertIn("78%", page_text)

    # ---------------------------------------------------------------------------
    # 4. FLASK ROUTE INTEGRATION TESTS
    # ---------------------------------------------------------------------------

    def test_export_version_resume_routes(self):
        v = create_resume_version(
            db_path=self.app.config["DATABASE_PATH"],
            resume_id=1,
            resume_details=self.sample_resume,
            filename="jane_v1.docx",
            file_path=Path(self.temp_dir.name) / "jane_v1.docx",
            template_filename="classic.docx",
        )

        # Test DOCX export
        res_docx = self.client.get(f"/export/resume/{v['id']}/docx")
        self.assertEqual(res_docx.status_code, 200)
        self.assertEqual(
            res_docx.mimetype,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment; filename=", res_docx.headers["Content-Disposition"])

        # Test PDF export
        res_pdf = self.client.get(f"/export/resume/{v['id']}/pdf")
        self.assertEqual(res_pdf.status_code, 200)
        self.assertEqual(res_pdf.mimetype, "application/pdf")
        self.assertTrue(res_pdf.data.startswith(b"%PDF"))

    def test_export_ats_report_route(self):
        # DOCX
        res_docx = self.client.post(
            "/export/ats-report/docx",
            data={"ats_data_json": json.dumps(self.sample_ats_data)},
        )
        self.assertEqual(res_docx.status_code, 200)
        self.assertEqual(
            res_docx.mimetype,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # PDF
        res_pdf = self.client.post(
            "/export/ats-report/pdf",
            data={"ats_data_json": json.dumps(self.sample_ats_data)},
        )
        self.assertEqual(res_pdf.status_code, 200)
        self.assertEqual(res_pdf.mimetype, "application/pdf")
        self.assertTrue(res_pdf.data.startswith(b"%PDF"))

    def test_export_jd_match_report_route(self):
        # DOCX
        res_docx = self.client.post(
            "/export/jd-match-report/docx",
            data={"match_data_json": json.dumps(self.sample_jd_match_data)},
        )
        self.assertEqual(res_docx.status_code, 200)
        self.assertEqual(
            res_docx.mimetype,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # PDF
        res_pdf = self.client.post(
            "/export/jd-match-report/pdf",
            data={"match_data_json": json.dumps(self.sample_jd_match_data)},
        )
        self.assertEqual(res_pdf.status_code, 200)
        self.assertEqual(res_pdf.mimetype, "application/pdf")
        self.assertTrue(res_pdf.data.startswith(b"%PDF"))

    def test_export_error_handling(self):
        # Missing ATS data
        res = self.client.post("/export/ats-report/pdf", data={})
        self.assertEqual(res.status_code, 302) # Redirects back with flash warning

        # Unsupported format
        res_unsupported = self.client.post(
            "/export/ats-report/txt",
            data={"ats_data_json": json.dumps(self.sample_ats_data)},
        )
        self.assertEqual(res_unsupported.status_code, 302)


if __name__ == "__main__":
    unittest.main()
