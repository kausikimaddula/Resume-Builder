import unittest
from pathlib import Path
from docx import Document
from services.resume_builder import (
    extract_template_outline,
    build_replacements,
    replace_template_placeholders,
)

class TestResumeBuilder(unittest.TestCase):
    def setUp(self):
        # Create a mock document in memory
        self.doc = Document()
        self.doc.add_paragraph("Hello, {{full_name}}!")
        self.doc.add_paragraph("Degree: {{degree}}")
        
        # Add a table
        table = self.doc.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = "Company: [company]"
        row.cells[1].text = "Role: [role]"
        
        # Details similar to routes
        self.resume_details = {
            "personal": {
                "full_name": "Jane User",
                "email": "jane@example.com",
                "phone": "555-0199",
                "linkedin": "linkedin.com/in/jane",
                "github": "github.com/jane",
                "portfolio": "jane.dev",
                "address": "123 Main St",
            },
            "education": {
                "degree": "B.S. Computer Science",
                "college": "Test University",
                "graduation_year": "2024",
                "gpa": "3.8",
            },
            "skills": "Python, Flask, WTForms",
            "experience": {
                "company": "Tech Corp",
                "role": "Software Engineer Intern",
                "duration": "3 months",
                "responsibilities": "Cleaned up unit tests.",
            },
            "projects": {
                "project_name": "AI Resume Builder",
                "description": "Built AI integration.",
                "technologies": "Python, OpenAI",
            },
            "certifications": "AWS Developer Associate",
            "achievements": "First place in hackathon",
            "languages": "English, Spanish",
        }

    def test_extract_template_outline(self):
        outline = extract_template_outline(self.doc)
        self.assertIn("paragraphs", outline)
        self.assertIn("tables", outline)
        self.assertIn("placeholders", outline)
        
        paragraphs = outline["paragraphs"]
        self.assertEqual(len(paragraphs), 2)
        self.assertEqual(paragraphs[0]["text"], "Hello, {{full_name}}!")
        
        tables = outline["tables"]
        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0][0], ["Company: [company]", "Role: [role]"])
        
        placeholders = outline["placeholders"]
        self.assertIn("{{full_name}}", placeholders)
        self.assertIn("{{degree}}", placeholders)
        self.assertIn("[company]", placeholders)
        self.assertIn("[role]", placeholders)

    def test_build_replacements(self):
        ai_response = {
            "placeholders": {
                "{{full_name}}": "Jane User",
                "{{degree}}": "B.S. Computer Science",
            },
            "resume_content": {}
        }
        
        replacements = build_replacements(self.resume_details, ai_response)
        
        # Verify LLM mappings
        self.assertEqual(replacements.get("{{full_name}}"), "Jane User")
        self.assertEqual(replacements.get("{{degree}}"), "B.S. Computer Science")
        
        # Verify static mappings expand properly
        self.assertEqual(replacements.get("{{company}}"), "Tech Corp")
        self.assertEqual(replacements.get("{company}"), "Tech Corp")
        self.assertEqual(replacements.get("[company]"), "Tech Corp")
        self.assertEqual(replacements.get("COMPANY"), "Tech Corp")

    def test_replace_template_placeholders(self):
        ai_response = {
            "placeholders": {
                "{{full_name}}": "Jane User",
                "{{degree}}": "B.S. Computer Science",
                "[company]": "Tech Corp",
                "[role]": "Software Engineer Intern",
            },
            "resume_content": {}
        }
        
        replaced_count = replace_template_placeholders(self.doc, self.resume_details, ai_response)
        self.assertTrue(replaced_count > 0)
        
        # Verify paragraph replacement
        self.assertEqual(self.doc.paragraphs[0].text, "Hello, Jane User!")
        self.assertEqual(self.doc.paragraphs[1].text, "Degree: B.S. Computer Science")
        
        # Verify table replacement
        self.assertEqual(self.doc.tables[0].rows[0].cells[0].text, "Company: Tech Corp")
        self.assertEqual(self.doc.tables[0].rows[0].cells[1].text, "Role: Software Engineer Intern")

if __name__ == "__main__":
    unittest.main()
