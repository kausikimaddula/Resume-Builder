import sys
from pathlib import Path
from docx import Document

def generate_docx():
    print("Generating DOCX...")
    doc = Document()
    doc.add_heading("Jane Doe - Resume", level=0)
    doc.add_paragraph("Email: jane.doe@example.com | Phone: (123) 456-7890")
    
    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Experienced software engineer specializing in Python development and technical writing.")
    
    # Skills table
    doc.add_heading("Skills", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Languages"
    hdr_cells[1].text = "Python, HTML, CSS, JavaScript"
    
    doc.save("test_resume.docx")
    print("Saved test_resume.docx")

def generate_pdf():
    print("Generating PDF...")
    text = (
        "John Smith - Resume\n"
        "Email: john.smith@example.com | Phone: +1-555-0199\n\n"
        "Summary\n"
        "Dedicated system administrator with background in cloud infrastructure design and deployment.\n\n"
        "Experience\n"
        "Cloud Engineer at Global Tech (2022-2025)\n"
        "Worked on AWS architectures and Python shell scripting.\n"
    )
    
    lines = text.split("\n")
    content_lines = ["BT", "/F1 12 Tf", "50 800 Td", "15 TL"]
    for l in lines:
        escaped = l.replace("(", "\\(").replace(")", "\\)")
        content_lines.append(f"({escaped}) Tj T*")
    content_lines.append("ET")
    content_stream = "\n".join(content_lines).encode("latin-1")
    
    objects = []
    def add_obj(data):
        objects.append(data)
        return len(objects)

    catalog_id = add_obj(b"<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_obj(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    page_id = add_obj(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>")
    contents_id = add_obj(f"<< /Length {len(content_stream)} >>\nstream\n{content_stream.decode('latin-1')}\nendstream".encode("latin-1"))
    font_id = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    
    with open("test_resume.pdf", "wb") as f:
        f.write(b"%PDF-1.4\n")
        offsets = []
        for i, obj in enumerate(objects):
            offsets.append(f.tell())
            f.write(f"{i+1} 0 obj\n".encode("latin-1"))
            f.write(obj)
            f.write(b"\nendobj\n")
            
        xref_offset = f.tell()
        f.write(b"xref\n")
        f.write(f"0 {len(objects)+1}\n".encode("latin-1"))
        f.write(b"0000000000 65535 f \n")
        for offset in offsets:
            f.write(f"{offset:010d} 00000 n \n".encode("latin-1"))
        f.write(b"trailer\n")
        f.write(f"<< /Size {len(objects)+1} /Root 1 0 R >>\n".encode("latin-1"))
        f.write(b"startxref\n")
        f.write(f"{xref_offset}\n".encode("latin-1"))
        f.write(b"%%EOF\n")
        
    print("Saved test_resume.pdf")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
