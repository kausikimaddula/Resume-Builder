import docx

doc = docx.Document()
# Add styled components
p1 = doc.add_paragraph()
p1.add_run("John Doe").bold = True
doc.add_paragraph("Email: john.doe@example.com")
doc.add_paragraph("Phone: {{phone}}")
doc.add_paragraph("LinkedIn: {{linkedin}}")

doc.add_heading("Education", level=1)
doc.add_paragraph("College: [college]")
doc.add_paragraph("Degree: {{degree}}")

doc.add_heading("Experience", level=1)
doc.add_paragraph("Company: [company]")
doc.add_paragraph("Role: [role]")
doc.add_paragraph("Responsibilities: {{responsibilities}}")

doc.save("c:/Users/SAI KAUSIKI/OneDrive/Desktop/Resume-Builder/uploads/test_template.docx")
print("Template created successfully!")
