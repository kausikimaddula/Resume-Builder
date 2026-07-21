"""Unit tests for services.job_description."""

from __future__ import annotations

import unittest
from pathlib import Path
from docx import Document
from werkzeug.datastructures import FileStorage
from io import BytesIO

from services.job_description import (
    extract_text_from_txt,
    extract_jd_text,
    save_jd_upload,
    UploadedJobDescription
)

class TestJobDescription(unittest.TestCase):
    """Tests for job description extraction and upload validation."""

    def setUp(self):
        self.tmp_dir = Path("tmp_test_uploads")
        self.tmp_dir.mkdir(exist_ok=True)
        self.created_files = []

    def tearDown(self):
        for file_path in self.created_files:
            if file_path.exists():
                file_path.unlink()
        if self.tmp_dir.exists():
            try:
                self.tmp_dir.rmdir()
            except OSError:
                pass

    def create_temp_file(self, filename: str, content: str | bytes) -> Path:
        file_path = self.tmp_dir / filename
        if isinstance(content, str):
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
        else:
            with open(file_path, "wb") as f:
                f.write(content)
        self.created_files.append(file_path)
        return file_path

    def test_extract_text_from_txt(self):
        """Test extracting text from TXT files."""
        expected_text = "Software Developer Job Description\nMust know Python & SQL."
        txt_path = self.create_temp_file("test_jd.txt", expected_text)
        
        extracted = extract_text_from_txt(txt_path)
        self.assertEqual(extracted, expected_text)

    def test_extract_jd_text_txt(self):
        """Test extract_jd_text with TXT files."""
        expected_text = "Require 3+ years experience."
        txt_path = self.create_temp_file("test_jd.txt", expected_text)
        
        extracted = extract_jd_text(txt_path)
        self.assertEqual(extracted, expected_text)

    def test_extract_jd_text_docx(self):
        """Test extract_jd_text with DOCX files."""
        doc = Document()
        doc.add_paragraph("We are hiring a Python Engineer.")
        p = doc.add_paragraph("Requirements:")
        p.add_run(" Flask experience is a plus.")
        
        docx_path = self.tmp_dir / "test_jd.docx"
        doc.save(docx_path)
        self.created_files.append(docx_path)

        extracted = extract_jd_text(docx_path)
        self.assertIn("We are hiring a Python Engineer.", extracted)
        self.assertIn("Flask experience is a plus.", extracted)

    def test_extract_jd_text_unsupported(self):
        """Test extract_jd_text raising ValueError on unsupported extension."""
        png_path = self.create_temp_file("test_image.png", "fake-binary-data")
        with self.assertRaises(ValueError):
            extract_jd_text(png_path)

    def test_save_jd_upload_success_txt(self):
        """Test saving a text job description file."""
        file_bytes = b"Need clean code principles."
        storage = FileStorage(
            stream=BytesIO(file_bytes),
            filename="jd_position.txt",
            content_type="text/plain",
        )
        
        uploaded = save_jd_upload(storage, self.tmp_dir)
        self.created_files.append(uploaded.path)

        self.assertEqual(uploaded.original_filename, "jd_position.txt")
        self.assertEqual(uploaded.extension, "txt")
        self.assertEqual(uploaded.file_type, "TXT Job Description")
        self.assertTrue(uploaded.path.exists())
        
        with open(uploaded.path, "r", encoding="utf-8") as f:
            self.assertEqual(f.read(), "Need clean code principles.")

    def test_save_jd_upload_rejected(self):
        """Test verifying disallowed file extensions represent ValueError."""
        storage = FileStorage(
            stream=BytesIO(b"fake pdf content"),
            filename="resume.png",
            content_type="image/png",
        )
        with self.assertRaises(ValueError):
            save_jd_upload(storage, self.tmp_dir)


if __name__ == "__main__":
    unittest.main()
