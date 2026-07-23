"""AI-powered DOCX resume generation service."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from shutil import copy2
from typing import Any
from uuid import uuid4

from docx import Document

# Import shared AI service helpers and centralized prompts
from services.openai_service import OpenAI, OpenAiServiceError, execute_json_chat_completion, parse_and_clean_json
from services.prompts import RESUME_BUILDER_SYSTEM_PROMPT


PLACEHOLDER_KEYS = {
    "full_name": ("personal", "full_name"),
    "email": ("personal", "email"),
    "phone": ("personal", "phone"),
    "linkedin": ("personal", "linkedin"),
    "github": ("personal", "github"),
    "portfolio": ("personal", "portfolio"),
    "address": ("personal", "address"),
    "degree": ("education", "degree"),
    "college": ("education", "college"),
    "graduation_year": ("education", "graduation_year"),
    "gpa": ("education", "gpa"),
    "skills": ("skills",),
    "company": ("experience", "company"),
    "role": ("experience", "role"),
    "duration": ("experience", "duration"),
    "responsibilities": ("experience", "responsibilities"),
    "project_name": ("projects", "project_name"),
    "project_description": ("projects", "description"),
    "technologies": ("projects", "technologies"),
    "certifications": ("certifications",),
    "achievements": ("achievements",),
    "languages": ("languages",),
}


@dataclass(frozen=True)
class GeneratedResume:
    """Details about the generated resume file."""

    filename: str
    path: Path


class ResumeBuilderError(Exception):
    """Raised when resume generation cannot be completed safely."""


def build_resume_from_template(
    *,
    resume_details: dict[str, Any],
    template_path: Path,
    output_folder: Path,
    api_key: str,
    model: str,
) -> GeneratedResume:
    """Generate a filled DOCX resume without modifying the original template."""
    if not api_key:
        raise ResumeBuilderError("OPENAI_API_KEY is missing. Add it to your .env file.")

    if template_path.suffix.lower() != ".docx":
        raise ResumeBuilderError("Only DOCX templates can be used for generation.")

    output_folder.mkdir(parents=True, exist_ok=True)
    generated_filename = f"completed-resume-{uuid4().hex[:8]}.docx"
    generated_path = output_folder / generated_filename

    # Copy first so all original DOCX styles, margins, headers, and formatting
    # remain intact. The uploaded template file itself is never changed.
    copy2(template_path, generated_path)

    try:
        document = Document(generated_path)
    except Exception as error:
        raise ResumeBuilderError(f"Failed to load the DOCX template: {error}") from error

    template_outline = extract_template_outline(document)
    
    ai_response = generate_resume_content(
        resume_details=resume_details,
        template_outline=template_outline,
        api_key=api_key,
        model=model,
    )

    replaced_count = replace_template_placeholders(document, resume_details, ai_response)
    
    # Fallback: if absolutely no placeholders or sample texts were matched, append content using default styling
    if replaced_count == 0:
        append_resume_content(document, ai_response.get("resume_content", {}))

    try:
        document.save(generated_path)
    except Exception as error:
        raise ResumeBuilderError(f"Failed to save the completed resume: {error}") from error

    return GeneratedResume(filename=generated_filename, path=generated_path)


def extract_template_outline(document: Document) -> dict[str, Any]:
    """Read visible DOCX text and style names so the model can infer structure."""
    paragraphs = [
        {"text": paragraph.text, "style": paragraph.style.name if paragraph.style else ""}
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    return {
        "paragraphs": paragraphs[:80],
        "tables": tables[:20],
        "placeholders": find_placeholders(document),
    }


def generate_resume_content(
    *,
    resume_details: dict[str, Any],
    template_outline: dict[str, Any],
    api_key: str,
    model: str,
) -> dict[str, Any]:
    """Ask OpenAI for dynamic placeholder mapping and polished resume content."""
    client = OpenAI(api_key=api_key)

    user_prompt_dict = {
        "task": "Build the mapping and polish the content to fit the template.",
        "user_details": resume_details,
        "template_outline": template_outline,
    }
    user_prompt = json.dumps(user_prompt_dict, ensure_ascii=True)

    try:
        raw_json = execute_json_chat_completion(
            system_prompt=RESUME_BUILDER_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            api_key=api_key,
            model=model,
            client=client,
        )
    except OpenAiServiceError as error:
        raise ResumeBuilderError(str(error)) from error

    return raw_json


def parse_json_response(raw_text: str) -> dict[str, Any]:
    """Parse model JSON, tolerating accidental fenced code blocks."""
    try:
        return parse_and_clean_json(raw_text)
    except OpenAiServiceError as error:
        raise ResumeBuilderError(str(error)) from error


def replace_template_placeholders(
    document: Document,
    resume_details: dict[str, Any],
    ai_response: dict[str, Any],
) -> int:
    """Replace placeholders and sample text in the template while preserving styles."""
    replacements = build_replacements(resume_details, ai_response)
    replaced_count = 0

    for paragraph in iter_all_paragraphs(document):
        replaced_count += replace_in_paragraph(paragraph, replacements)

    return replaced_count


def build_replacements(
    resume_details: dict[str, Any],
    ai_response: dict[str, Any],
) -> dict[str, str]:
    """Create placeholder values from user data and AI mapping."""
    replacements: dict[str, str] = {}

    # 1. Start with placeholders mapped by the LLM
    ai_placeholders = ai_response.get("placeholders", {})
    if isinstance(ai_placeholders, dict):
        for key, value in ai_placeholders.items():
            replacements[str(key)] = stringify(value)

    # 2. Build static mappings for standard keys to guarantee coverage
    static_vals = {}
    for key, path in PLACEHOLDER_KEYS.items():
        value = nested_get(resume_details, path)
        static_vals[key] = stringify(value)

    # Incorporate dynamic content from the AI response if present
    ai_content = ai_response.get("resume_content", {})
    if isinstance(ai_content, dict):
        static_vals["headline"] = stringify(ai_content.get("headline"))
        static_vals["contact_line"] = stringify(ai_content.get("contact_line"))

        # Merge arrays to strings
        for field in ["education", "skills", "certifications", "achievements", "languages"]:
            if field in ai_content:
                static_vals[field] = stringify(ai_content.get(field))

    # Expand standard keys into common formats (e.g. {{full_name}}, {full_name}, [full_name], FULL_NAME)
    for key, value in static_vals.items():
        replacements[f"{{{{{key}}}}}"] = value
        replacements[f"{{{key}}}"] = value
        replacements[f"[{key}]"] = value
        replacements[key.upper()] = value

    # Filter out empty keys
    return {k: v for k, v in replacements.items() if k.strip()}


def replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> int:
    """Replace placeholder values in a paragraph, matching longer patterns first to avoid fragments."""
    count = 0
    sorted_keys = sorted(replacements.keys(), key=len, reverse=True)

    for key in sorted_keys:
        replacement_value = replacements[key]
        if key in paragraph.text:
            # Try to replace within single runs first (most precise, keeps formatting)
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, replacement_value)
                    count += 1

            # Fallback for keys split across multiple runs
            loop_guard = 0
            while key in paragraph.text and loop_guard < 10:
                loop_guard += 1
                success = replace_text_across_runs(paragraph, key, replacement_value)
                if not success:
                    break
                count += 1

    return count


def replace_text_across_runs(paragraph: Any, search_text: str, replacement_text: str) -> bool:
    """Replace search_text across multiple runs in a paragraph while preserving styling."""
    text = paragraph.text
    start_idx = text.find(search_text)
    if start_idx == -1:
        return False
    end_idx = start_idx + len(search_text)

    current_pos = 0
    run_ranges = []
    for run in paragraph.runs:
        run_len = len(run.text)
        run_ranges.append((current_pos, current_pos + run_len))
        current_pos += run_len

    overlapping_runs = []
    for idx, (r_start, r_end) in enumerate(run_ranges):
        if max(start_idx, r_start) < min(end_idx, r_end):
            overlapping_runs.append(idx)

    if not overlapping_runs:
        return False

    first_run_idx = overlapping_runs[0]
    last_run_idx = overlapping_runs[-1]

    first_run = paragraph.runs[first_run_idx]
    last_run = paragraph.runs[last_run_idx]

    first_run_start, _ = run_ranges[first_run_idx]
    last_run_start, _ = run_ranges[last_run_idx]

    prefix = first_run.text[:(start_idx - first_run_start)]
    suffix = last_run.text[(end_idx - last_run_start):]

    if first_run_idx == last_run_idx:
        first_run.text = prefix + replacement_text + suffix
    else:
        first_run.text = prefix + replacement_text
        for idx in overlapping_runs[1:-1]:
            paragraph.runs[idx].text = ""
        last_run.text = suffix

    return True


def append_resume_content(document: Document, resume_content: dict[str, Any]) -> None:
    """Fallback: append the resume content to the end of the document."""
    document.add_heading(stringify(resume_content.get("full_name")), level=1)
    add_optional_paragraph(document, resume_content.get("headline"))
    add_optional_paragraph(document, resume_content.get("contact_line"))

    add_list_section(document, "Education", resume_content.get("education", []))
    add_list_section(document, "Skills", resume_content.get("skills", []))

    experiences = resume_content.get("experience", [])
    if experiences:
        document.add_heading("Experience", level=2)
        for item in experiences:
            title = " - ".join(
                part
                for part in [
                    stringify(item.get("title")),
                    stringify(item.get("company")),
                    stringify(item.get("duration")),
                ]
                if part
            )
            add_optional_paragraph(document, title)
            for bullet in item.get("bullets", []):
                document.add_paragraph(stringify(bullet), style="List Bullet")

    projects = resume_content.get("projects", [])
    if projects:
        document.add_heading("Projects", level=2)
        for item in projects:
            add_optional_paragraph(document, stringify(item.get("name")))
            add_optional_paragraph(document, stringify(item.get("description")))
            add_optional_paragraph(document, f"Technologies: {stringify(item.get('technologies'))}")

    add_list_section(document, "Certifications", resume_content.get("certifications", []))
    add_list_section(document, "Achievements", resume_content.get("achievements", []))
    add_list_section(document, "Languages", resume_content.get("languages", []))


def add_list_section(document: Document, title: str, items: list[Any]) -> None:
    """Append a heading and bullet list when content exists."""
    clean_items = [stringify(item) for item in items if stringify(item)]
    if not clean_items:
        return

    document.add_heading(title, level=2)
    for item in clean_items:
        document.add_paragraph(item, style="List Bullet")


def add_optional_paragraph(document: Document, value: Any) -> None:
    """Append one paragraph when the value is not empty."""
    text = stringify(value)
    if text:
        document.add_paragraph(text)


def iter_all_paragraphs(document: Document):
    """Yield paragraphs from the body, tables, headers, and footers."""
    yield from document.paragraphs

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs

    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def find_placeholders(document: Document) -> list[str]:
    """Find common placeholder tokens in the template text."""
    tokens: set[str] = set()
    pattern = re.compile(r"(\{\{[^{}]+\}\}|\{[^{}]+\}|\[[^\[\]]+\])")
    for paragraph in iter_all_paragraphs(document):
        tokens.update(pattern.findall(paragraph.text))
    return sorted(tokens)


def nested_get(data: dict[str, Any], path: tuple[str, ...]) -> Any:
    """Read a nested value from a dictionary."""
    current: Any = data
    for key in path:
        if not isinstance(current, dict):
            return ""
        current = current.get(key, "")
    return current


def stringify(value: Any) -> str:
    """Convert simple values and lists to display text."""
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(stringify(item) for item in value)
    return str(value).strip()
