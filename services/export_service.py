"""Service for exporting Generated Resumes, ATS Reports, and JD Match Reports to DOCX and PDF formats."""

from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable

logger = logging.getLogger(__name__)


class ExportServiceError(Exception):
    """Raised when document export (DOCX/PDF) fails."""


# ---------------------------------------------------------------------------
# RESUME EXPORTERS
# ---------------------------------------------------------------------------

def export_resume_docx(resume_details: dict[str, Any]) -> bytes:
    """Generate a clean, styled DOCX document from resume details dictionary."""
    try:
        doc = Document()
        
        # Set standard margins (0.75 in)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        personal = resume_details.get("personal", {})
        full_name = personal.get("full_name", "Resume")
        
        # Title / Name Header
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(full_name)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        title_p.paragraph_format.space_after = Pt(4)

        # Contact Info Line
        contact_parts = []
        if personal.get("email"):
            contact_parts.append(personal["email"])
        if personal.get("phone"):
            contact_parts.append(personal["phone"])
        if personal.get("linkedin"):
            contact_parts.append(f"LinkedIn: {personal['linkedin']}")
        if personal.get("github"):
            contact_parts.append(f"GitHub: {personal['github']}")
        if personal.get("portfolio"):
            contact_parts.append(f"Portfolio: {personal['portfolio']}")
        if personal.get("address"):
            contact_parts.append(personal["address"])

        if contact_parts:
            contact_p = doc.add_paragraph(" | ".join(contact_parts))
            contact_p.paragraph_format.space_after = Pt(14)
            for r in contact_p.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

        # Helper function for section headings
        def _add_section_heading(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title.upper())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        # Education
        edu = resume_details.get("education", {})
        if isinstance(edu, dict) and any(edu.values()):
            _add_section_heading("Education")
            deg = edu.get("degree", "")
            col = edu.get("college", "")
            yr = edu.get("graduation_year", "")
            gpa = edu.get("gpa", "")

            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_after = Pt(4)
            r_deg = edu_p.add_run(deg)
            r_deg.font.bold = True
            if col:
                edu_p.add_run(f" — {col}")
            if yr:
                edu_p.add_run(f" ({yr})")
            if gpa:
                edu_p.add_run(f" | GPA: {gpa}")

        # Experience
        exp = resume_details.get("experience", {})
        if isinstance(exp, dict) and any(exp.values()):
            _add_section_heading("Experience")
            role = exp.get("role", "")
            comp = exp.get("company", "")
            dur = exp.get("duration", "")
            resp = exp.get("responsibilities", "")

            exp_p = doc.add_paragraph()
            exp_p.paragraph_format.space_after = Pt(2)
            r_role = exp_p.add_run(role)
            r_role.font.bold = True
            if comp:
                exp_p.add_run(f" at {comp}")
            if dur:
                exp_p.add_run(f" ({dur})")

            if resp:
                for line in resp.splitlines():
                    line_str = line.strip()
                    if line_str:
                        if line_str.startswith("-") or line_str.startswith("•"):
                            line_str = line_str.lstrip("-• ").strip()
                        bp = doc.add_paragraph(line_str, style="List Bullet")
                        bp.paragraph_format.space_after = Pt(2)

        # Projects
        proj = resume_details.get("projects", {})
        if isinstance(proj, dict) and any(proj.values()):
            _add_section_heading("Projects")
            pname = proj.get("project_name", "")
            pdesc = proj.get("description", "")
            ptech = proj.get("technologies", "")

            proj_p = doc.add_paragraph()
            proj_p.paragraph_format.space_after = Pt(2)
            r_pname = proj_p.add_run(pname)
            r_pname.font.bold = True
            if ptech:
                proj_p.add_run(f" ({ptech})")

            if pdesc:
                for line in pdesc.splitlines():
                    line_str = line.strip()
                    if line_str:
                        if line_str.startswith("-") or line_str.startswith("•"):
                            line_str = line_str.lstrip("-• ").strip()
                        bp = doc.add_paragraph(line_str, style="List Bullet")
                        bp.paragraph_format.space_after = Pt(2)

        # Skills
        if resume_details.get("skills"):
            _add_section_heading("Skills")
            sp = doc.add_paragraph(str(resume_details["skills"]))
            sp.paragraph_format.space_after = Pt(4)

        # Certifications
        if resume_details.get("certifications"):
            _add_section_heading("Certifications")
            cp = doc.add_paragraph(str(resume_details["certifications"]))
            cp.paragraph_format.space_after = Pt(4)

        # Achievements
        if resume_details.get("achievements"):
            _add_section_heading("Achievements")
            ap = doc.add_paragraph(str(resume_details["achievements"]))
            ap.paragraph_format.space_after = Pt(4)

        # Languages
        if resume_details.get("languages"):
            _add_section_heading("Languages")
            lp = doc.add_paragraph(str(resume_details["languages"]))
            lp.paragraph_format.space_after = Pt(4)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as error:
        logger.exception("Failed to export resume DOCX: %s", error)
        raise ExportServiceError(f"Could not generate DOCX resume: {error}") from error


def export_resume_pdf(resume_details: dict[str, Any]) -> bytes:
    """Generate a clean, styled PDF document from resume details using ReportLab."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            "ResumeName",
            parent=styles["Heading1"],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#111827"),
            spaceAfter=4,
            fontName="Helvetica-Bold",
        )
        
        contact_style = ParagraphStyle(
            "ResumeContact",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=12,
            fontName="Helvetica",
        )

        heading_style = ParagraphStyle(
            "ResumeSectionHeading",
            parent=styles["Heading2"],
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F2937"),
            spaceBefore=10,
            spaceAfter=4,
            fontName="Helvetica-Bold",
        )

        body_style = ParagraphStyle(
            "ResumeBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=4,
            fontName="Helvetica",
        )

        bullet_style = ParagraphStyle(
            "ResumeBullet",
            parent=body_style,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=3,
        )

        story = []

        personal = resume_details.get("personal", {})
        full_name = personal.get("full_name", "Resume")
        story.append(Paragraph(full_name, name_style))

        contact_parts = []
        if personal.get("email"):
            contact_parts.append(personal["email"])
        if personal.get("phone"):
            contact_parts.append(personal["phone"])
        if personal.get("linkedin"):
            contact_parts.append(f"LinkedIn: {personal['linkedin']}")
        if personal.get("github"):
            contact_parts.append(f"GitHub: {personal['github']}")
        if personal.get("portfolio"):
            contact_parts.append(f"Portfolio: {personal['portfolio']}")
        if personal.get("address"):
            contact_parts.append(personal["address"])

        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), contact_style))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E5E7EB"), spaceAfter=10))

        # Education
        edu = resume_details.get("education", {})
        if isinstance(edu, dict) and any(edu.values()):
            story.append(Paragraph("EDUCATION", heading_style))
            deg = edu.get("degree", "")
            col = edu.get("college", "")
            yr = edu.get("graduation_year", "")
            gpa = edu.get("gpa", "")

            edu_line = f"<b>{deg}</b>"
            if col:
                edu_line += f" — {col}"
            if yr:
                edu_line += f" ({yr})"
            if gpa:
                edu_line += f" | GPA: {gpa}"
            story.append(Paragraph(edu_line, body_style))

        # Experience
        exp = resume_details.get("experience", {})
        if isinstance(exp, dict) and any(exp.values()):
            story.append(Paragraph("EXPERIENCE", heading_style))
            role = exp.get("role", "")
            comp = exp.get("company", "")
            dur = exp.get("duration", "")
            resp = exp.get("responsibilities", "")

            exp_line = f"<b>{role}</b>"
            if comp:
                exp_line += f" at {comp}"
            if dur:
                exp_line += f" ({dur})"
            story.append(Paragraph(exp_line, body_style))

            if resp:
                for line in resp.splitlines():
                    line_str = line.strip()
                    if line_str:
                        if line_str.startswith("-") or line_str.startswith("•"):
                            line_str = line_str.lstrip("-• ").strip()
                        story.append(Paragraph(f"• {line_str}", bullet_style))

        # Projects
        proj = resume_details.get("projects", {})
        if isinstance(proj, dict) and any(proj.values()):
            story.append(Paragraph("PROJECTS", heading_style))
            pname = proj.get("project_name", "")
            pdesc = proj.get("description", "")
            ptech = proj.get("technologies", "")

            proj_line = f"<b>{pname}</b>"
            if ptech:
                proj_line += f" ({ptech})"
            story.append(Paragraph(proj_line, body_style))

            if pdesc:
                for line in pdesc.splitlines():
                    line_str = line.strip()
                    if line_str:
                        if line_str.startswith("-") or line_str.startswith("•"):
                            line_str = line_str.lstrip("-• ").strip()
                        story.append(Paragraph(f"• {line_str}", bullet_style))

        # Skills
        if resume_details.get("skills"):
            story.append(Paragraph("SKILLS", heading_style))
            story.append(Paragraph(str(resume_details["skills"]), body_style))

        # Certifications
        if resume_details.get("certifications"):
            story.append(Paragraph("CERTIFICATIONS", heading_style))
            story.append(Paragraph(str(resume_details["certifications"]), body_style))

        # Achievements
        if resume_details.get("achievements"):
            story.append(Paragraph("ACHIEVEMENTS", heading_style))
            story.append(Paragraph(str(resume_details["achievements"]), body_style))

        # Languages
        if resume_details.get("languages"):
            story.append(Paragraph("LANGUAGES", heading_style))
            story.append(Paragraph(str(resume_details["languages"]), body_style))

        doc.build(story)
        return buffer.getvalue()
    except Exception as error:
        logger.exception("Failed to export resume PDF: %s", error)
        raise ExportServiceError(f"Could not generate PDF resume: {error}") from error


# ---------------------------------------------------------------------------
# ATS REPORT EXPORTERS
# ---------------------------------------------------------------------------

def export_ats_report_docx(ats_data: dict[str, Any]) -> bytes:
    """Generate ATS score report in DOCX format."""
    try:
        doc = Document()
        
        # Margins
        for s in doc.sections:
            s.top_margin = Inches(0.75)
            s.bottom_margin = Inches(0.75)
            s.left_margin = Inches(0.75)
            s.right_margin = Inches(0.75)

        # Title
        title_p = doc.add_paragraph()
        r_title = title_p.add_run("ATS Compatibility Analysis Report")
        r_title.font.size = Pt(20)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep blue

        # Subtitle / Score Box
        score = ats_data.get("score", 0)
        atype = ats_data.get("analysis_type", "ATS Evaluation")
        
        score_p = doc.add_paragraph()
        score_p.paragraph_format.space_before = Pt(6)
        score_p.paragraph_format.space_after = Pt(12)
        
        r_score_lbl = score_p.add_run("Overall ATS Compatibility Score: ")
        r_score_lbl.font.size = Pt(13)
        
        r_score_val = score_p.add_run(f"{score} / 100")
        r_score_val.font.size = Pt(16)
        r_score_val.font.bold = True
        if score >= 75:
            r_score_val.font.color.rgb = RGBColor(0x16, 0x65, 0x34) # Green
        elif score >= 50:
            r_score_val.font.color.rgb = RGBColor(0xCA, 0x8A, 0x04) # Yellow
        else:
            r_score_val.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) # Red

        score_p.add_run(f"  ({atype})")

        # Section Helper
        def _add_section(title: str, items: list[str], color_rgb: RGBColor):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title)
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = color_rgb

            if items:
                for item in items:
                    bp = doc.add_paragraph(str(item), style="List Bullet")
                    bp.paragraph_format.space_after = Pt(3)
            else:
                doc.add_paragraph("None identified.")

        _add_section("Key Strengths", ats_data.get("strengths", []), RGBColor(0x16, 0x65, 0x34))
        _add_section("Areas for Improvement", ats_data.get("weaknesses", []), RGBColor(0xDC, 0x26, 0x26))
        _add_section("Actionable Recommendations", ats_data.get("suggestions", []), RGBColor(0x1D, 0x4E, 0x89))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as error:
        logger.exception("Failed to export ATS report DOCX: %s", error)
        raise ExportServiceError(f"Could not generate ATS report DOCX: {error}") from error


def export_ats_report_pdf(ats_data: dict[str, Any]) -> bytes:
    """Generate ATS score report in PDF format using ReportLab."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "ATSTitle",
            parent=styles["Heading1"],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#1E3A8A"),
            fontName="Helvetica-Bold",
            spaceAfter=6,
        )

        score_style = ParagraphStyle(
            "ATSScore",
            parent=styles["Normal"],
            fontSize=14,
            leading=18,
            fontName="Helvetica-Bold",
            spaceAfter=14,
        )

        section_heading = ParagraphStyle(
            "ATSSectionHeading",
            parent=styles["Heading2"],
            fontSize=13,
            leading=16,
            fontName="Helvetica-Bold",
            spaceBefore=12,
            spaceAfter=6,
        )

        bullet_style = ParagraphStyle(
            "ATSBullet",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=4,
            fontName="Helvetica",
        )

        story = []

        story.append(Paragraph("ATS Compatibility Analysis Report", title_style))

        score = ats_data.get("score", 0)
        atype = ats_data.get("analysis_type", "ATS Evaluation")
        
        score_color = "#166534" if score >= 75 else ("#CA8A04" if score >= 50 else "#DC2626")
        score_text = f"Overall ATS Score: <font color='{score_color}'>{score} / 100</font> <font size=10 color='#6B7280'>({atype})</font>"
        story.append(Paragraph(score_text, score_style))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E5E7EB"), spaceAfter=10))

        # Strengths
        story.append(Paragraph("<font color='#166534'>Key Strengths</font>", section_heading))
        strengths = ats_data.get("strengths", [])
        if strengths:
            for item in strengths:
                story.append(Paragraph(f"• {item}", bullet_style))
        else:
            story.append(Paragraph("None identified.", bullet_style))

        # Weaknesses
        story.append(Paragraph("<font color='#DC2626'>Areas for Improvement</font>", section_heading))
        weaknesses = ats_data.get("weaknesses", [])
        if weaknesses:
            for item in weaknesses:
                story.append(Paragraph(f"• {item}", bullet_style))
        else:
            story.append(Paragraph("None identified.", bullet_style))

        # Suggestions
        story.append(Paragraph("<font color='#1D4E89'>Actionable Recommendations</font>", section_heading))
        suggestions = ats_data.get("suggestions", [])
        if suggestions:
            for item in suggestions:
                story.append(Paragraph(f"• {item}", bullet_style))
        else:
            story.append(Paragraph("None identified.", bullet_style))

        doc.build(story)
        return buffer.getvalue()
    except Exception as error:
        logger.exception("Failed to export ATS report PDF: %s", error)
        raise ExportServiceError(f"Could not generate ATS report PDF: {error}") from error


# ---------------------------------------------------------------------------
# JD MATCH REPORT EXPORTERS
# ---------------------------------------------------------------------------

def export_jd_match_report_docx(match_data: dict[str, Any]) -> bytes:
    """Generate Resume vs Job Description Match report in DOCX format."""
    try:
        doc = Document()

        for s in doc.sections:
            s.top_margin = Inches(0.75)
            s.bottom_margin = Inches(0.75)
            s.left_margin = Inches(0.75)
            s.right_margin = Inches(0.75)

        # Header
        title_p = doc.add_paragraph()
        r_title = title_p.add_run("Job Description Match Analysis Report")
        r_title.font.size = Pt(20)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark slate

        match_pct = match_data.get("match_percentage", 0)
        atype = match_data.get("analysis_type", "JD Matcher")

        score_p = doc.add_paragraph()
        score_p.paragraph_format.space_after = Pt(12)
        score_p.add_run("Job Description Match Score: ").font.size = Pt(13)
        r_val = score_p.add_run(f"{match_pct}%")
        r_val.font.size = Pt(16)
        r_val.font.bold = True
        r_val.font.color.rgb = RGBColor(0x16, 0x65, 0x34) if match_pct >= 75 else (
            RGBColor(0x02, 0x84, 0xC7) if match_pct >= 50 else RGBColor(0xCA, 0x8A, 0x04)
        )
        score_p.add_run(f"  ({atype})")

        def _add_section(title: str, items: list[str], color_rgb: RGBColor = RGBColor(0x1E, 0x29, 0x3B)):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(title)
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = color_rgb

            if items:
                for item in items:
                    bp = doc.add_paragraph(str(item), style="List Bullet")
                    bp.paragraph_format.space_after = Pt(2)
            else:
                doc.add_paragraph("None identified.")

        _add_section("Matching Skills & Qualifications", match_data.get("matching_skills", []), RGBColor(0x16, 0x65, 0x34))
        _add_section("Missing Technical Skills", match_data.get("missing_technical_skills", []), RGBColor(0xDC, 0x26, 0x26))
        _add_section("Missing Soft Skills", match_data.get("missing_soft_skills", []), RGBColor(0xD9, 0x77, 0x06))
        _add_section("Recommended Keywords to Add", match_data.get("recommended_keywords", []), RGBColor(0x25, 0x63, 0xEB))
        _add_section("Recommended Certifications", match_data.get("recommended_certifications", []), RGBColor(0x7C, 0x3A, 0xED))
        _add_section("Suggested Projects to Prove Skills", match_data.get("recommended_projects", []), RGBColor(0x0D, 0x94, 0x88))
        _add_section("Step-by-Step Action Roadmap", match_data.get("learning_roadmap", []), RGBColor(0x0F, 0x17, 0x2A))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as error:
        logger.exception("Failed to export JD match report DOCX: %s", error)
        raise ExportServiceError(f"Could not generate JD match report DOCX: {error}") from error


def export_jd_match_report_pdf(match_data: dict[str, Any]) -> bytes:
    """Generate Resume vs Job Description Match report in PDF format using ReportLab."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "JDTitle",
            parent=styles["Heading1"],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0F172A"),
            fontName="Helvetica-Bold",
            spaceAfter=6,
        )

        score_style = ParagraphStyle(
            "JDScore",
            parent=styles["Normal"],
            fontSize=14,
            leading=18,
            fontName="Helvetica-Bold",
            spaceAfter=14,
        )

        section_heading = ParagraphStyle(
            "JDSectionHeading",
            parent=styles["Heading2"],
            fontSize=12,
            leading=15,
            fontName="Helvetica-Bold",
            spaceBefore=10,
            spaceAfter=4,
        )

        bullet_style = ParagraphStyle(
            "JDBullet",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13.5,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=3,
            fontName="Helvetica",
        )

        story = []

        story.append(Paragraph("Job Description Match Analysis Report", title_style))

        match_pct = match_data.get("match_percentage", 0)
        atype = match_data.get("analysis_type", "JD Matcher")
        
        score_color = "#166534" if match_pct >= 75 else ("#0284C7" if match_pct >= 50 else "#CA8A04")
        score_text = f"Job Match Score: <font color='{score_color}'>{match_pct}%</font> <font size=10 color='#6B7280'>({atype})</font>"
        story.append(Paragraph(score_text, score_style))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E5E7EB"), spaceAfter=8))

        def _add_pdf_section(title: str, items: list[str], color_hex: str):
            story.append(Paragraph(f"<font color='{color_hex}'>{title}</font>", section_heading))
            if items:
                for item in items:
                    story.append(Paragraph(f"• {item}", bullet_style))
            else:
                story.append(Paragraph("None identified.", bullet_style))

        _add_pdf_section("Matching Skills & Qualifications", match_data.get("matching_skills", []), "#166534")
        _add_pdf_section("Missing Technical Skills", match_data.get("missing_technical_skills", []), "#DC2626")
        _add_pdf_section("Missing Soft Skills", match_data.get("missing_soft_skills", []), "#D97706")
        _add_pdf_section("Recommended Keywords to Add", match_data.get("recommended_keywords", []), "#2563EB")
        _add_pdf_section("Recommended Certifications", match_data.get("recommended_certifications", []), "#7C3AED")
        _add_pdf_section("Suggested Projects to Prove Skills", match_data.get("recommended_projects", []), "#0D9488")
        _add_pdf_section("Step-by-Step Action Roadmap", match_data.get("learning_roadmap", []), "#0F172A")

        doc.build(story)
        return buffer.getvalue()
    except Exception as error:
        logger.exception("Failed to export JD match report PDF: %s", error)
        raise ExportServiceError(f"Could not generate JD match report PDF: {error}") from error
