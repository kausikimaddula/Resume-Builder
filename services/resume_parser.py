"""Service for parsing and extracting text from PDF and DOCX resume files."""

from __future__ import annotations

import logging
from pathlib import Path
from docx import Document
from pypdf import PdfReader

from services.exceptions import InvalidFileError

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract all text page-by-page from a PDF file."""
    logger.info("Extracting text from PDF: %s", file_path)
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for index, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                logger.warning("Empty page text retrieved from page %d of %s", index, file_path)
        
        extracted = "\n".join(text_parts).strip()
        logger.info("Extracted %d characters from PDF: %s", len(extracted), file_path)
        return extracted
    except Exception as e:
        logger.error("Failed to extract text from PDF file '%s': %s", file_path, e, exc_info=True)
        raise InvalidFileError(
            message=f"Error reading PDF file {file_path}: {e}",
            user_message="Could not read text from the uploaded PDF file. The file may be damaged or password-protected.",
        ) from e


def extract_text_from_docx(file_path: Path) -> str:
    """Extract all text from paragraphs and tables in a DOCX file."""
    logger.info("Extracting text from DOCX: %s", file_path)
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Read tables to not miss crucial content like resume grids
        for table in doc.tables:
            for row in table.rows:
                # Merge duplicate cells in rows (often happens with merged columns)
                cells_text = []
                for cell in row.cells:
                    c_text = cell.text.strip()
                    if c_text and (not cells_text or cells_text[-1] != c_text):
                        cells_text.append(c_text)
                if cells_text:
                    text_parts.append(" | ".join(cells_text))
                    
        extracted = "\n".join(text_parts).strip()
        logger.info("Extracted %d characters from DOCX: %s", len(extracted), file_path)
        return extracted
    except Exception as e:
        logger.error("Failed to extract text from DOCX file '%s': %s", file_path, e, exc_info=True)
        raise InvalidFileError(
            message=f"Error reading DOCX file {file_path}: {e}",
            user_message="Could not read text from the uploaded DOCX file. The document may be corrupted.",
        ) from e


def extract_resume_text(file_path: Path) -> str:
    """Detect file type and extract text from the resume.

    Supported formats: PDF, DOCX.
    """
    if not file_path.exists():
        logger.error("Resume file not found at path: %s", file_path)
        raise InvalidFileError(
            message=f"Resume file not found at: {file_path}",
            user_message="The uploaded resume file could not be found on the server.",
        )

    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    else:
        logger.warning("Unsupported file type '%s' requested for text extraction.", suffix)
        raise InvalidFileError(
            message=f"Unsupported resume file type: {suffix}",
            user_message=f"Unsupported file format '{suffix}'. Only PDF and DOCX files are allowed.",
        )
