"""Temporary test script for services.resume_parser."""

import sys
from pathlib import Path
from docx import Document
from services.resume_parser import extract_resume_text

def test_docx_parser():
    test_file = Path("test_resume.docx")
    
    # Create test DOCX
    doc = Document()
    doc.add_paragraph("Jane Doe Resume")
    doc.add_paragraph("Software Engineer")
    
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Project A'
    hdr_cells[1].text = 'Python, Flask, React'
    
    doc.save(test_file)
    print("Created test DOCX file.")

    try:
        text = extract_resume_text(test_file)
        print("--- Extracted Text ---")
        print(text)
        print("----------------------")
        assert "Jane Doe Resume" in text
        assert "Software Engineer" in text
        assert "Project A" in text
        assert "Python, Flask" in text
        print("DOCX extraction test passed!")
    finally:
        if test_file.exists():
            test_file.unlink()

if __name__ == "__main__":
    try:
        test_docx_parser()
    except Exception as e:
        print(f"Test failed: {e}")
        sys.exit(1)
